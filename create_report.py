"""生成详细汇报大纲Word文档"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ─── Styles ───
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = '微软雅黑'
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.3

for level in range(1, 4):
    h = doc.styles[f'Heading {level}']
    h.font.name = '微软雅黑'
    h.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

# ═══════════════════════════════════════════
# Cover
# ═══════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('学历信息AI自动核验系统')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('演示汇报大纲')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x55, 0x6B, 0x8D)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f'日期: {datetime.date.today().strftime("%Y年%m月%d日")}').font.size = Pt(12)
info.add_run('\n技术栈: Python + PaddleOCR + DeepSeek V4 Pro + Streamlit').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════
# Table of Contents
# ═══════════════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目背景与痛点分析',
    '二、整体逻辑流程（五层架构）',
    '三、核心功能详细介绍',
    '  3.1 上传层',
    '  3.2 识别层',
    '  3.3 核验层 — 模块A: 单证单项核验',
    '  3.4 核验层 — 模块B: 多材料交叉验证',
    '  3.5 判定分流层',
    '  3.6 人工核验台',
    '  3.7 数据看板',
    '四、技术方案选型与关键决策',
    '五、特殊情况处理矩阵',
    '  5.1 证书层面特殊场景',
    '  5.2 简历层面特殊场景',
    '  5.3 招聘类型差异',
    '  5.4 业务边界与HR流程',
    '六、Demo演示安排',
    '七、后续扩展方向',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════
# 一、项目背景
# ═══════════════════════════════════════════
doc.add_heading('一、项目背景与痛点分析', level=1)

doc.add_heading('1.1 业务场景', level=2)
doc.add_paragraph('每年校招季，HR部门需处理大量候选人的学历材料。每位候选人至少提供学位证书和毕业证书各一份（本/硕/博可达6份），HR需要逐一核对证书上的姓名、学校、专业、学历层次、毕业日期等关键信息，并与简历进行交叉比对。')

doc.add_heading('1.2 传统人工核验的痛点', level=2)
pain_points = [
    ('耗时长', '批量校招期间，手动核对数百份证书需要数天甚至数周时间，严重拖慢招聘进度。'),
    ('易遗漏', '长时间重复工作导致注意力下降，细微的信息差异（如专业名称、日期格式）容易被忽略。'),
    ('造假风险', '证书图片可通过PS篡改姓名、日期等关键字段。人工肉眼难以识别精修过的图片。'),
    ('知识盲区', 'HR无法记住全国3000+所高校、数万个专业的所有组合。医科大学是否开设计算机专业？某所学校是否改过名？这些判断超出了HR的知识范围。'),
    ('无系统记录', '核验结果散落在Excel和邮件中，无法追溯、统计和审计。'),
]
for item in pain_points:
    title, desc = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('1.3 项目目标', level=2)
goals = [
    '构建一套AI驱动的学历信息自动验证系统，覆盖证书识别、关键字段提取、多源数据交叉比对、异常预警全流程',
    '将HR的单人核验时间从数分钟压缩到秒级，批量处理能力从数十份提升到数百份',
    '通过多维检测（教育部名单、学信库、ELA篡改检测、双证比对、简历交叉验证）降低漏判率',
    '提供人工核验台和数据看板，实现核验结果的可追溯、可统计、可审计',
]
for g in goals:
    doc.add_paragraph(g, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 二、整体逻辑流程
# ═══════════════════════════════════════════
doc.add_heading('二、整体逻辑流程（五层架构）', level=1)

doc.add_heading('2.1 五层架构总览', level=2)
layers = [
    ('上传层', '输入层', '校招/社招/实习类型选择 → 单人/批量上传模式 → 学位证+毕业证+简历（必传）'),
    ('识别层', 'AI解析层', 'PaddleOCR文字识别（印章去除+增强） → DeepSeek V4 Pro结构化提取（证书7字段+简历多段教育经历）'),
    ('核验层', '风控引擎', '模块A: 单证单项核验（6项检查） + 模块B: 多材料交叉验证（双证比对+简历分层匹配）'),
    ('判定分流层', '结果输出', 'PASS自动通过 / REVIEW需复核 / ALERT告警 → 人工核验台HR最终判定'),
    ('数据看板', '统计层', '分类名单（通过/不通过/待定/待核验）+ 内联修改状态 + CSV导出'),
]
for item in layers:
    name, role, desc = item[0], item[1], item[2]
    p = doc.add_paragraph()
    run = p.add_run(f'{name}（{role}）：')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.2 用户操作三步流', level=2)
steps = [
    'Step 1/2: 选择招聘场景（校招/社招/实习） → 选择上传模式（单人/批量） → 上传学位证+毕业证+简历',
    'Step 2/2: 系统自动完成OCR识别+LLM提取+多维核验 → 展示结果：总览表+逐层详情+简历交叉验证',
]
for s in steps:
    doc.add_paragraph(s, style='List Number')

doc.add_heading('2.3 核验引擎内部流程', level=2)
doc.add_paragraph('对于每份上传的证书，系统依次执行以下检查：')
checks_flow = [
    '教育部名单校验 → 学校是否在583所国内本科+191所海外认证名单中？曾用名自动解析。',
    '学信库查询 → 按身份证号/证书编号/姓名三级fallback查询学籍记录，逐字段比对（姓名/学校/专业/学历/日期）。',
    '证书状态检测 → 识别结业证书、肄业证书。',
    '毕业年龄检测 → 根据学历层次判断年龄合理性（本科18-30岁，硕士22-40岁，博士24-45岁）。',
    '院校-专业一致性 → 正向查找该学校是否真实开设该专业（基于12,697条教育部数据）。',
    'ELA图片篡改检测 → JPEG压缩误差分析，检测图片是否经过PS修改（三级判定: low/medium/high）。',
]
for i, c in enumerate(checks_flow, 1):
    doc.add_paragraph(f'{i}. {c}')

doc.add_paragraph('以上6项为模块A（单证单项核验）。完成后进入模块B：')
cross_checks = [
    '双证交叉比对 → 学位证vs毕业证的姓名/学校/专业/学历/日期逐字段比对。',
    '简历分层交叉验证 → 按学历层级逐段与对应证书比对：学校（简称/曾用名归一化）+ 专业（包含匹配）+ 年份（严格匹配）。',
]
for c in cross_checks:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 三、核心功能
# ═══════════════════════════════════════════
doc.add_heading('三、核心功能详细介绍', level=1)

# 3.1 上传层
doc.add_heading('3.1 上传层', level=2)
doc.add_paragraph('上传层承担所有输入入口，支持灵活的场景适配。')

upload_features = [
    ('招聘类型选择', '校招/社招/实习三选一。选择后系统自动切换判定规则：校招缺证→REVIEW待取证，社招缺证→ALERT不通过，实习已毕业缺证→ALERT、在读→跳过。'),
    ('单人上传', '动态添加学历层级（本科/硕士/博士），每层级上传2份证书。学位证为必传，毕业证可选（同等学力申硕、部分海外国家仅发学位证）。简历为必传项。'),
    ('批量上传', '上传ZIP压缩包，每个子文件夹对应一位候选人。系统自动检测顶层包装文件夹并跳过。按文件名关键词（本科/硕士/博士/学位/毕业/简历）自动分类证书类型。'),
    ('文件格式支持', 'JPG / PNG / PDF。PDF通过pypdfium2自动转换为图片后再进行OCR。上传的证书原件自动保存至本地目录，供人工核验台查看。'),
]
for item in upload_features:
    title, desc = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

# 3.2 识别层
doc.add_heading('3.2 识别层', level=2)

doc.add_heading('3.2.1 OCR文字识别', level=3)
ocr_desc = [
    '引擎: PaddleOCR 2.6.2，本地离线运行，无需联网，无调用次数限制。',
    '预处理: (1) HSV色彩空间分离红色印章区域并替换为背景色；(2) 自适应缩放至2000px以内；(3) CLAHE对比度增强 + 非局部均值降噪。',
    '双通道识别: 对去印章图和原图分别做OCR，合并结果并去重排序，兼顾印章覆盖区域的文字和正常区域的文字。',
    '参数优化: 检测阈值降至0.2（默认0.3），适配证书细字体；开启角度分类处理倾斜拍照。',
]
for d in ocr_desc:
    doc.add_paragraph(d, style='List Bullet')

doc.add_heading('3.2.2 AI字段提取', level=3)
llm_desc = [
    '引擎: DeepSeek V4 Pro（Anthropic兼容API），通过自然语言理解提取结构化字段，零正则规则。',
    '证书提取字段: 姓名、性别、出生日期、学校、学院、专业、学历层次、学位、学习形式、毕业日期、入学日期、证书编号、证书类型、签发日期。共14个字段。',
    '简历提取字段: 多段教育经历数组，每段含学校、专业、学历层次、入学年份、毕业年份、状态标签(graduated/ongoing/dropout/exchange/minor/merged)。',
    '日期智能区分: 通过上下文关键词区分出生日期（"XX年XX月生"）、毕业日期（"于XX年毕业"）、签发日期（底部签名日期）。学位证无独立毕业日期时自动取签发日期。',
    '中英文键名映射: AI可能返回中文键名（如"姓名""学校"），通过KEY_MAP自动转换为英文字段名。',
    '学位证vs毕业证区分: 先识别证书类型，学位证不提取身份证号和学习形式字段。',
]
for d in llm_desc:
    doc.add_paragraph(d, style='List Bullet')

# 3.3 核验层A
doc.add_heading('3.3 核验层 — 模块A: 单证单项核验', level=2)
doc.add_paragraph('每份证书独立执行的6项检查，每项直接输出判定级别（PASS/REVIEW/ALERT）：')

# Table
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['检查项', 'PASS（通过）', 'REVIEW（需复核）', 'ALERT（告警）']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

rules = [
    ('教育部名单', '学校在583国内+191海外名单中', '—', '学校不在名单中，直接拒绝'),
    ('学信库查询', '查到记录且字段一致', '查到但部分字段不一致', '未查到该学籍记录'),
    ('证书状态', '毕业/学位证书', '结业/肄业证书', '—'),
    ('毕业年龄', '年龄在正常范围内', '偏大或偏小', '—'),
    ('院校-专业一致性', '专业在学校公开列表中', '专业未在列表中', '—'),
    ('ELA图片篡改', '未检测到(risk<0.30)', '轻微异常(0.30-0.40)', '明显篡改(>0.40)'),
]
for i, (name, p, r, a) in enumerate(rules):
    row = table.rows[i+1]
    row.cells[0].text = name
    row.cells[1].text = p
    row.cells[2].text = r
    row.cells[3].text = a

# 3.4 核验层B
doc.add_heading('3.4 核验层 — 模块B: 多材料交叉验证', level=2)

doc.add_heading('3.4.1 双证交叉比对', level=3)
doc.add_paragraph('同一学历层级的学位证与毕业证，逐字段比对：姓名、学校、专业、学历层次、毕业日期。任一项不一致→标记REVIEW。若毕业证未上传（同等学力/海外），该层级只做学位证核验，双证比对自动跳过。')

doc.add_heading('3.4.2 简历分层交叉验证', level=3)
doc.add_paragraph('核心原则：将简历提取的多段教育经历，按学历层级分别与对应证书组进行比对，不跨层。')

resume_rules = [
    ('状态预处理', 'LLM提取每段经历的status标签。ongoing(在读)/dropout(退学)/exchange(交换)/minor(辅修)/merged(连读)各有不同的处理策略。'),
    ('学校匹配', '证书学校与简历学校比较。支持简全称归一化（"北大"→"北京大学"）和曾用名解析（"北京钢铁学院"→"北京科技大学"）。'),
    ('专业匹配', '采用包含关系匹配。证书专业"法学"，简历写"民法学"→法学⊆民法学→✅ 通过。证书专业"历史学"，简历写"计算机科学"→完全不包含→❌ 不通过。'),
    ('年份匹配', '证书毕业年份与简历end_year严格相等。不一致→标记REVIEW（可能是延期/提前毕业）。'),
    ('缺证处理', '简历声称已毕业但未上传对应证书：社招/实习→ALERT；校招→REVIEW"应届待取证"。校招中在读(ongoing)的学历也标记为REVIEW"应届待取证"。'),
]
for item in resume_rules:
    title, desc = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

# 3.5 判定分流
doc.add_heading('3.5 判定分流层', level=2)
doc.add_paragraph('综合所有检查项的结果，按优先级汇总：')
doc.add_paragraph('存在至少一项ALERT → 最终判定ALERT（严重异常）', style='List Bullet')
doc.add_paragraph('无ALERT，存在至少一项REVIEW → 最终判定REVIEW（需复核）', style='List Bullet')
doc.add_paragraph('全部检查均为PASS → 最终判定PASS（全部通过）', style='List Bullet')

doc.add_paragraph('')
doc.add_paragraph('PASS的候选人自动进入通过名单，无需经过人工核验台。')
doc.add_paragraph('ALERT和REVIEW的候选人进入人工核验台，HR查看材料和系统判定原因后，做出最终决定。')

# 3.6 人工核验台
doc.add_heading('3.6 人工核验台', level=2)
doc.add_paragraph('专为HR设计的终审界面，分为两栏：')

review_features = [
    ('待核验栏', '展示所有系统判定为REVIEW或ALERT、尚未经HR判定的候选人。每个候选人展开后显示：系统判定级别（ALERT/REVIEW，颜色区分）、详细的未通过原因列表（仅列出非PASS项）、候选人原始证书材料（按类型分组折叠：学位证书/毕业证书/简历/HR补充材料，点击可查看大图）。'),
    ('HR判定', '三选一：PASS通过、FAIL淘汰、HOLD待补充材料。附带审核备注栏。判定后候选人自动移入已核验栏。'),
    ('已核验栏', '展示HR已判定过的候选人。保留完整材料查看功能，支持随时重新判定（如候选人后续补交了证书，HR可将HOLD改为PASS）。'),
    ('补充材料上传', 'HR可为候选人手动补传文件（如成绩单、学校证明、学信网验证报告等），文件自动存入该候选人材料文件夹的"HR补充"子目录。'),
]
for item in review_features:
    title, desc = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}：')
    run.bold = True
    p.add_run(desc)

# 3.7 数据看板
doc.add_heading('3.7 数据看板', level=2)
doc.add_paragraph('汇总全部候选人数据的统计视图：')

dashboard = [
    '分类标签页：全部候选人 / 通过 / 不通过 / 待定 / 待核验（五个标签），每个标签内展示对应候选人列表。',
    '内联状态修改：在表格中直接通过下拉框修改候选人状态，修改后自动归类到对应分组，无需额外操作。',
    '显示字段：姓名、学校、学历、招聘类型（校招/社招/实习）、HR判定状态、审核备注、更新时间。',
    '导出功能：每个标签页支持导出CSV文件，方便HR进行线下汇总和报告。',
    '数据基础：覆盖583所国内本科院校、191所教育部认证海外高校、51条院校曾用名映射、12,697条院校-专业数据。',
]
for d in dashboard:
    doc.add_paragraph(d, style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、技术方案
# ═══════════════════════════════════════════
doc.add_heading('四、技术方案选型与关键决策', level=1)

doc.add_heading('4.1 技术栈总览', level=2)
tech_table = doc.add_table(rows=9, cols=4)
tech_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['层级', '技术选型', '版本/型号', '选型理由']):
    tech_table.rows[0].cells[i].text = h
    for p in tech_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

techs = [
    ('OCR引擎', 'PaddleOCR', '2.6.2', '本地离线运行，中文识别最优（证书为中文文档），免费无调用限制。2.x版本在Windows上稳定，3.x存在PIR兼容问题。'),
    ('AI字段提取', 'DeepSeek V4 Pro', '—', '通过Anthropic兼容API调用。零样本理解证书文字，无需标注训练。自然语言理解能力远强于正则表达式，能处理各种证书格式差异。已配置API Key和Base URL。'),
    ('图片预处理', 'OpenCV', '4.11', 'HSV印章去除、自适应缩放、CLAHE增强、非局部均值降噪，均为OpenCV内置函数。'),
    ('图片篡改检测', 'OpenCV (ELA)', '—', '误差水平分析(Error Level Analysis)，通过对比原图与重新JPEG压缩后的差异来检测PS痕迹。纯算法实现，无需额外模型。'),
    ('前端UI', 'Streamlit', '1.59', 'Python全栈框架，无需学习HTML/CSS/JS即可构建交互式Web应用。内置文件上传、数据表格、图片显示等组件。30天交付期内最快方案。'),
    ('数据库', 'SQLite', '—', '单机Demo零配置。数据量在演示范围内（<1000候选人），性能完全足够。生产环境可迁移至PostgreSQL。'),
    ('PDF处理', 'pypdfium2', '—', '将PDF文件渲染为PNG图片供OCR处理，核验台中将PDF简历渲染为图片供HR查看。比PyMuPDF更轻量。'),
    ('PPT生成', 'python-pptx', '—', '自动化生成演示PPT，支持深色主题、圆角矩形、自定义配色。'),
]
for i, (layer, tech, version, reason) in enumerate(techs):
    row = tech_table.rows[i+1]
    row.cells[0].text = layer
    row.cells[1].text = tech
    row.cells[2].text = version
    row.cells[3].text = reason

doc.add_heading('4.2 关键架构决策', level=2)
decisions = [
    ('为何用LLM而非正则提取字段？', '证书格式千差万别。同样是学位证，不同学校的文字表述完全不同："学科门类为管理学"、"在XX专业修完教学计划规定的全部课程"、"已通过课程考试和论文答辩"。正则规则永远追不上格式的多样性。LLM能理解自然语言，无论什么表述都能准确提取。DeepSeek V4 Pro通过Anthropic兼容接口调用，零训练成本。'),
    ('为何用正向查找而非关键词否定？', '曾经将"医科大学"与"计算机科学与技术"视为冲突。但实际搜索证实，徐州医科大学、山东第一医科大学等确实开设了计算机专业。正确的做法是维护一个可更新的院校-专业数据库，正向查找"该专业是否在学校列表"而非"该关键词是否被禁止"。'),
    ('为何PaddleOCR版本锁定2.6.2？', 'PaddlePaddle 3.x引入了PIR（Paddle IR），与Windows系统上的oneDNN不兼容，反复报错。2.x版本使用传统执行路径，稳定可靠。3.x新增的API变化（predict()→ocr()）也不兼容现有代码。'),
    ('为何学信库和教育部名单分开？', '教育部名单验证"学校是否真实存在"，学信库查询"该学生是否有学籍记录"。这是两个独立的概念：一所真实存在的大学里，可能没有名为"张三"的毕业生。两步独立核验才能精准定位问题。'),
    ('为何使用Streamlit而非前后端分离？', '项目交付周期仅一个月，且需要面向HR演示。Streamlit可以在Python文件中同时处理UI和逻辑，无需管理API接口、前端框架、状态同步等复杂度。Demo阶段效率优先。'),
]
for item in decisions:
    title, desc = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}\n')
    run.bold = True
    p.add_run(desc)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════════
# 五、特殊情况
# ═══════════════════════════════════════════
doc.add_heading('五、特殊情况处理矩阵', level=1)

doc.add_heading('5.1 证书层面特殊场景', level=2)
cert_cases = [
    ('红色印章遮挡文字', 'HSV色彩空间分离红色区域→替换为背景色→OCR识别', '学位证书的圆形红色印章是最常见的OCR干扰源，去除后可显著提升识别率'),
    ('高分辨率大图', '自适应缩放至2000px以内', '原始证书照片可达5000-6000px，直接OCR会导致文字检测失败。2000px是PaddleOCR的最佳识别范围'),
    ('细字体漏检', '检测阈值由默认0.3降至0.2', '证书上的印刷字体可能比自然场景文字更细，默认阈值会漏掉'),
    ('学校曾用名/更名', '曾用名映射表（51条）自动解析为现名', '如简历写"北京钢铁学院"，证书写"北京科技大学"。映射表覆盖常见高校更名历史'),
    ('结业/肄业冒充毕业', '证书类型识别，结业肄业→标记REVIEW', '结业证书和毕业证书外观相似，但有本质区别'),
    ('图片PS篡改', 'ELA三级判定：low→PASS / medium→REVIEW / high→ALERT', '通过对比原图与JPEG重压缩图的误差分布，检测修改过的区域。阈值经过多次校准（low<0.30, medium 0.30-0.40, high>0.40）'),
    ('无毕业证的情况', '毕业证改为可选上传', '同等学力申硕只有学位证无毕业证；部分海外国家仅发学位证。该场景下只验证学位证，双证比对自动跳过'),
    ('学位证与毕业证格式差异', '双通道OCR + LLM自动识别证书类型', '学位证不提取身份证号和学习形式字段；毕业证可提取这些字段'),
]
for item in cert_cases:
    title, solution, note = item[0], item[1], item[2]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.bold = True
    p.add_run(f'\n处理方式: {solution}\n说明: {note}')

doc.add_heading('5.2 简历层面特殊场景', level=2)
resume_cases = [
    ('专业名称细化', '简历写"民法学"，证书写"法学"。采用包含关系匹配：法学⊆民法学→通过。专业完全不同→不通过', '防止正常细化被误判，同时抓住"计算机科学"冒充"会计学"的造假'),
    ('学校简称', '简历写"北大"，证书写"北京大学"。通过字符串包含和归一化比较匹配', '高校简称在简历中非常普遍'),
    ('延期/提前毕业', '年份严格匹配。不一致→标记REVIEW，不阻断', '延期毕业不等于造假，但需要HR确认原因'),
    ('在读/未毕业', '简历写"至今"或"present"或end_year≥当前年份→status=ongoing。校招→标记"应届待取证"，不阻断；社招→不存在此情况', 'LLM自动识别文字信号和年份逻辑双重判定'),
    ('退学/肄业', 'status=dropout→自动跳过该段教育经历，不要求证书', '不能因为退学经历而认为候选人造假'),
    ('交换/访学', 'status=exchange→自动跳过', '交换经历通常无该校学位证，不应要求证书'),
    ('辅修/双学位', 'status=minor→自动跳过', '辅修可能无独立证书，或只有辅修证明而非学位证'),
    ('硕博/本博连读', 'status=merged→合并为一段，只验最高学位', '八年制临床医学等项目中，本硕博阶段不独立，只有一份最终证书'),
    ('直博无硕士证', '博士已毕业但无硕士证书→标记"可能为直博/硕博连读"，不阻断', '越来越多的博士项目接受本科直申，跳过硕士阶段是合法的学术路径'),
    ('多段同层次学历', '每段独立核验，每段都要通过', '如双本科（转学/第二学士学位），不能因为一段通过就漏掉另一段'),
    ('简历写了学位但未上传该层级证书', '校招→REVIEW"应届待取证"；社招→ALERT"缺证"；实习（已毕业）→ALERT"缺证"', '校招允许"待取证"是因为应届生尚未拿到证书；社招不允许因为候选人早已毕业'),
]
for item in resume_cases:
    title, solution = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.bold = True
    p.add_run(f'\n{solution}')

doc.add_heading('5.3 招聘类型判定差异', level=2)
rec_table = doc.add_table(rows=5, cols=4)
rec_table.style = 'Light Grid Accent 1'
for i, h in enumerate(['场景', '校招', '社招', '实习']):
    rec_table.rows[0].cells[i].text = h
    for p in rec_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.font.bold = True

rec_data = [
    ('已毕业学历，有证书', '正常核验', '正常核验', '正常核验'),
    ('已毕业学历，缺证书', 'REVIEW 应届待取证', 'ALERT 缺证不通过', 'ALERT 缺证不通过'),
    ('在读中（ongoing）', 'REVIEW 应届待取证', '不存在（社招无在读）', '跳过，不要求证书'),
    ('dropout/exchange/minor', '跳过', '跳过', '跳过'),
]
for i, (scene, school, social, intern) in enumerate(rec_data):
    row = rec_table.rows[i+1]
    row.cells[0].text = scene
    row.cells[1].text = school
    row.cells[2].text = social
    row.cells[3].text = intern

doc.add_heading('5.4 业务边界与HR流程', level=2)
biz_cases = [
    ('校招生缺证的判定', 'HOLD（待补充材料），而非PASS。成绩单、学校证明等替代材料不能等同于学历证书。证书是唯一的法定学历凭证。PASS意味着"已核验完毕可以录用"，而缺证意味着"尚未完成核验"。HOLD状态明确提示HR后续需要追要证书。补证后HR在已核验栏重新判定。'),
    ('实习的定位', '实习生在读学历（ongoing）跳过不验，但已毕业学历必须验证。实习转正时HR重新走校招流程，不继承实习期间的判定结果。'),
    ('HR补充材料', '人工核验台提供文件上传功能，HR可手动补传成绩单、学校证明、学信网验证报告等。文件自动存入对应候选人材料文件夹的"HR补充"子目录。'),
    ('已判定可重新修改', '已核验栏中的候选人支持随时修改判定。例如：候选人补交了证书后，HR可将HOLD改为PASS。修改即刻生效，候选人自动重新归类。'),
    ('同名候选人', '系统通过UUID绑定每次上传会话，不依赖姓名做唯一标识。即使两位候选人都叫"张伟"，他们的材料和核验结果不会混淆。'),
    ('简历必传', '简历上传已改为必传项。未上传简历的系统报错提示，无法开始核验。简历是交叉验证的必要输入，没有简历无法进行分层匹配。'),
]
for item in biz_cases:
    title, desc = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.bold = True
    p.add_run(f'\n{desc}')

doc.add_page_break()

# ═══════════════════════════════════════════
# 六、Demo演示
# ═══════════════════════════════════════════
doc.add_heading('六、Demo演示安排', level=1)

doc.add_heading('6.1 演示场景', level=2)
scenarios = [
    ('场景一：单人校招核验（正常证书）', '上传张文宇的学位证+毕业证+简历（中国人民大学，劳动经济学）。展示完整流程：Step 1选择校招→Step 2上传材料→核验结果。预期结果：PASS通过，所有检查项无异常。'),
    ('场景二：图片篡改检测', '上传真证+假证。假证为程序生成的带噪点版本，ELA检测到篡改→ALERT告警。真证正常通过。展示ELA三级判定的效果。'),
    ('场景三：批量核验', '上传包含两位候选人的ZIP压缩包（张文宇+赵亮）。展示批量处理：逐人进度条→汇总结果表格→逐人可展开详情→CSV导出。'),
    ('场景四：人工核验台', '展示待核验栏中的REVIEW/ALERT候选人，演示HR查看材料、填写备注、三选一判定（通过/淘汰/待补充）的全流程。'),
    ('场景五：数据看板', '展示分类名单、内联状态修改、CSV导出功能。'),
]
for item in scenarios:
    title, desc = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(f'{title}')
    run.bold = True
    p.add_run(f'\n{desc}')

doc.add_heading('6.2 测试数据', level=2)
doc.add_paragraph('学信库已录入2条真实记录：')
doc.add_paragraph('张文宇 | 中国人民大学 | 劳动经济学 | 本科 | 1000242024000969', style='List Bullet')
doc.add_paragraph('赵亮 | 湘潭大学 | 法学 | 本科 | 105301202405000692', style='List Bullet')

doc.add_page_break()

# ═══════════════════════════════════════════
# 七、后续扩展
# ═══════════════════════════════════════════
doc.add_heading('七、后续扩展方向', level=1)

extensions = [
    ('对接真实数据源', [
        '对接学信网官方API，替代模拟学信库',
        '对接教育部留学服务中心（zwfw.cscse.edu.cn），实现海外学历自动认证',
        '院校-专业数据定时从教育部官网同步最新备案结果',
    ]),
    ('增强检测能力', [
        '接入更专业的图片篡改检测模型（如Noiseprint、ManTra-Net）替代基础ELA',
        '增加证书模板匹配：以学校为维度建立证书视觉特征库，检测格式异常',
        '支持证书二维码/条形码扫描验证（部分新版本证书已内置）',
    ]),
    ('系统集成', [
        '企业微信/飞书/钉钉通知：核验结果自动推送至HR',
        '对接ATS（招聘管理系统）：核验结果自动同步至候选人档案',
        '单点登录（SSO）集成企业账号体系',
    ]),
    ('数据分析', [
        '历史核验数据统计分析：造假率、常见造假手段、高风险院校排名',
        '异常模式挖掘：同一学校、同一批次出现多个异常证书时自动预警',
    ]),
]
for item in extensions:
    title, items = item[0], item[1]
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ─── Footer ───
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run('— 文档结束 —').font.size = Pt(10)
p.add_run(f'\n生成日期: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}').font.size = Pt(9)

doc.save('D:/edu-verify/学历核验系统演示汇报大纲.docx')
print('Saved: 学历核验系统演示汇报大纲.docx')
